#!/usr/bin/env python3
"""Convert docx paper to markdown and extract tables as JSON."""

from docx import Document
from docx.oxml.ns import qn
import json
import re

DOCX_PATH = "/home/lulu444/Downloads/熊羿成 上海外国语大学 认知语义学视角下具身智能基座模型的“具身性潜能”初探.docx"

doc = Document(DOCX_PATH)


def get_paragraph_style(para):
    """Detect heading level or list style."""
    style_name = para.style.name if para.style else "Normal"
    if style_name.startswith("Heading"):
        level = style_name.replace("Heading ", "")
        try:
            return f"h{int(level)}"
        except ValueError:
            pass
    # Detect Chinese numbered headings like "一、" "二、"
    text = para.text.strip()
    if re.match(r'^[一二三四五六七八九十]+、', text):
        return "h2"
    if re.match(r'^\d+\.\d+', text):
        return "h3"
    if re.match(r'^\d+\.\d+\.\d+', text):
        return "h4"
    return None


def para_to_markdown(para):
    """Convert a paragraph to markdown string."""
    text = para.text.strip()
    if not text:
        return ""

    style = get_paragraph_style(para)
    if style:
        prefix = "#" * int(style[1])
        return f"{prefix} {text}"

    # Handle bold/italic runs
    parts = []
    for run in para.runs:
        t = run.text
        if not t:
            continue
        if run.bold and run.italic:
            parts.append(f"***{t}***")
        elif run.bold:
            parts.append(f"**{t}**")
        elif run.italic:
            parts.append(f"*{t}*")
        else:
            parts.append(t)

    result = "".join(parts) if parts else text

    # Detect RQ lines
    if text.startswith("RQ"):
        return f"**{text}**"

    return result


def table_to_markdown(table):
    """Convert a table to markdown format."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(cells)

    if not rows:
        return ""

    # Build markdown table
    col_count = max(len(r) for r in rows)
    # Normalize column count
    for r in rows:
        while len(r) < col_count:
            r.append("")

    lines = []
    lines.append("| " + " | ".join(rows[0]) + " |")
    lines.append("| " + " | ".join(["---"] * col_count) + " |")
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)


def table_to_json(table, table_index):
    """Convert a table to structured JSON."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(cells)

    if not rows:
        return {}

    # Table 0: VLA测评研究总结
    if table_index == 0:
        header = rows[0]
        data = []
        for row in rows[1:]:
            entry = {}
            for i, h in enumerate(header):
                val = row[i] if i < len(row) else ""
                entry[h] = val
            data.append(entry)
        return {
            "table_id": "table_1_vla_evaluation_studies",
            "title": "表1 VLA测评研究语言理解能力考察总结",
            "description": "对先前18篇VLA测评研究的语言理解能力考察情况总结",
            "columns": header,
            "data": data
        }

    # Table 1: 丢类动词描述性特征
    if table_index == 1:
        # Skip the title row (row 0) and header rows (row 1, 2)
        # Row 1 is the main header, Row 2 is the sub-header
        main_header = rows[1]  # ['', '力度 (FORCE)', '手部高度 (HAND)', ...]
        sub_header = rows[2]   # ['', '均值 (SD)', '均值 (SD)', ...]

        # Build combined headers
        combined_headers = []
        for i in range(len(main_header)):
            mh = main_header[i]
            sh = sub_header[i] if i < len(sub_header) else ""
            if mh and sh and mh != sh:
                combined_headers.append(f"{mh} - {sh}")
            elif mh:
                combined_headers.append(mh)
            else:
                combined_headers.append(sh)

        # Parse Chinese and English sections
        chinese_verbs = []
        english_verbs = []
        current_section = None

        for row in rows[3:]:
            verb_name = row[0]
            if "汉语" in verb_name:
                current_section = "chinese"
                continue
            if "英语" in verb_name:
                current_section = "english"
                continue
            if not verb_name:
                continue

            entry = {
                "verb": verb_name,
                "force_mean_sd": row[1] if len(row) > 1 else "",
                "hand_height_mean_sd": row[2] if len(row) > 2 else "",
                "arm_bent_straight": row[3] if len(row) > 3 else "",
                "horizontal_direction": row[4] if len(row) > 4 else "",
                "vertical_trajectory": row[5] if len(row) > 5 else ""
            }

            if current_section == "chinese":
                chinese_verbs.append(entry)
            elif current_section == "english":
                english_verbs.append(entry)

        return {
            "table_id": "table_2_throw_verb_characteristics",
            "title": "表2 丢类动词的描述性特征（Gao, 2016）",
            "description": "汉语和英语投掷类动词在五个运动维度上的描述性统计数据",
            "dimensions": {
                "force": "力度，五点李克特量表，5=极强，1=极轻",
                "hand_height": "手部高度，以脚部地面为0，头顶为10的相对高度",
                "arm_posture": "手臂姿态，弯曲:伸直的频率比",
                "horizontal_direction": "水平方向，向前:向侧的频率比",
                "vertical_trajectory": "垂直轨迹，向上:向下的频率比"
            },
            "chinese_verbs": chinese_verbs,
            "english_verbs": english_verbs
        }


# === Build markdown ===
md_lines = []
md_lines.append('# 认知语义学视角下具身智能基座模型的“具身性”潜能初探')
md_lines.append("")
md_lines.append("**熊羿成** 上海外国语大学")
md_lines.append("")
md_lines.append("**李檬希** 西南大学")
md_lines.append("")
md_lines.append("**关键词：** 大语言模型 视觉语言模型 具身智能 认知语言学 人工智能测评")
md_lines.append("")

table_index = 0
para_index = 0

for element in doc.element.body:
    tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

    if tag == "p":
        # Find the corresponding paragraph object
        para = doc.paragraphs[para_index]
        para_index += 1
        md = para_to_markdown(para)
        if md:
            md_lines.append(md)
            md_lines.append("")
    elif tag == "tbl":
        # Find the corresponding table object
        table = doc.tables[table_index]
        md_lines.append(f"**[表{table_index + 1}]**")
        md_lines.append("")
        md_lines.append(table_to_markdown(table))
        md_lines.append("")
        table_index += 1

markdown_content = "\n".join(md_lines)

# === Write markdown ===
with open("/home/lulu444/Enactment-experiment/paper.md", "w", encoding="utf-8") as f:
    f.write(markdown_content)

# === Extract tables as JSON ===
tables_json = []
for i in range(len(doc.tables)):
    tables_json.append(table_to_json(doc.tables[i], i))

with open("/home/lulu444/Enactment-experiment/tables.json", "w", encoding="utf-8") as f:
    json.dump(tables_json, f, ensure_ascii=False, indent=2)

print("Done!")
print(f"Markdown: paper.md ({len(markdown_content)} chars)")
print(f"Tables JSON: tables.json ({len(tables_json)} tables)")
